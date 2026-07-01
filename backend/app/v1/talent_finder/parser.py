"""
Resume Parser Service
Handles PDF and DOCX parsing with text normalization
Preserves important characters like dashes for date range detection
"""

import fitz  # PyMuPDF
from docx import Document
import re
import io


class ResumeParser:
    """Deterministic resume parser supporting PDF and DOCX formats"""

    def parse_pdf(self, content: bytes) -> str:
        """Extract text from PDF bytes using PyMuPDF"""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            raw_text = "\n".join(text_parts)
            return self._normalize_text(raw_text)
        except Exception as e:
            print(f"PDF parsing error: {e}")
            return ""

    def parse_docx(self, content: bytes) -> str:
        """Extract text from DOCX bytes using python-docx"""
        try:
            doc = Document(io.BytesIO(content))
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            raw_text = "\n".join(text_parts)
            return self._normalize_text(raw_text)
        except Exception as e:
            print(f"DOCX parsing error: {e}")
            return ""

    def _normalize_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        IMPORTANT: Preserves dashes (converted to hyphen) for date range detection.
        """
        # Step 1: Convert all Unicode dash variants to regular hyphen BEFORE stripping
        # This is critical for date range detection (e.g., "2021 – Present")
        text = re.sub(r'[\u2013\u2014\u2012\u2015\u2212]', '-', text)  # en-dash, em-dash, etc.

        # Step 2: Convert bullet points to newlines for better section detection
        text = re.sub(r'[\u2022\u2023\u25cf\u25cb\u25aa\u25e6\u2043]', '\n', text)

        # Step 3: Remove other non-printable/non-ASCII characters
        # Keep: printable ASCII (0x20-0x7E), newlines, tabs, and hyphens
        text = re.sub(r'[^\x20-\x7E\n\t]', ' ', text)

        # Step 4: Normalize whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Step 5: Strip each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        text = text.strip()

        return text
